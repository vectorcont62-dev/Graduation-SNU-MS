from pathlib import Path
import re
import zipfile

from docx import Document
from pypdf import PdfReader


ROOT = Path(r"C:\Users\joshs\Desktop\local-laptop\Graduation-SNU-MS\석사졸업논문")
DOCX = ROOT / ".thesis_build" / "석사학위논문_진명_초안_working.docx"
PDF = ROOT / ".thesis_build" / "석사학위논문_진명_초안_working.pdf"

document = Document(DOCX)
styles = [paragraph.style.name for paragraph in document.paragraphs]
full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

print("sections", len(document.sections))
print("inline_shapes", len(document.inline_shapes))
print(
    "heading1",
    styles.count("Heading 1"),
    "heading2",
    styles.count("Heading 2"),
    "references",
    styles.count("References"),
)
print(
    "old_terms",
    {term: term in full_text for term in ("수직전하재분포", "이 찬")},
)
print("title_ok", "극저온 3차원 낸드 플래시 메모리의" in full_text)

reader = PdfReader(str(PDF))
pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
print("pdf_pages", len(reader.pages))
print(
    "error_phrases",
    {
        phrase: phrase in pdf_text
        for phrase in ("필드 업데이트 필요", "목차 항목을 찾을 수 없습니다", "Error!")
    },
)
print(
    "pdf_old_terms",
    {term: term in pdf_text for term in ("수직전하재분포", "이 찬")},
)

with zipfile.ZipFile(DOCX) as archive:
    media = [name for name in archive.namelist() if name.startswith("word/media/")]
    relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    document_xml = archive.read("word/document.xml").decode("utf-8")

image_relationship_ids = re.findall(
    r'<Relationship[^>]+Id="([^"]+)"[^>]+'
    r'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"',
    relationships,
)
print("package_media", len(media))
print(
    "main_image_rels",
    len(image_relationship_ids),
    "used_main_image_rels",
    sum(relationship_id in document_xml for relationship_id in image_relationship_ids),
)
