from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

BUILD_DIR = Path(__file__).resolve().parent
ROOT = BUILD_DIR.parent
REFERENCE = ROOT / "석사학위논문_이찬_초안.docx"
TARGET_COVER = ROOT / "석사학위논문_진명_초안.docx"
FIGURE_DIR = BUILD_DIR / "figures"
OUTPUT = BUILD_DIR / "석사학위논문_진명_초안_working.docx"

sys.path.insert(0, str(BUILD_DIR))
from thesis_content import (  # noqa: E402
    CHAPTERS,
    ENGLISH_ABSTRACT,
    ENGLISH_TITLE_LINES,
    KOREAN_ABSTRACT,
    KOREAN_TITLE_LINES,
    REFERENCES,
)


def clear_body_keep_final_sectpr(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def insert_body_element(document: Document, element) -> None:
    body = document._element.body
    body.insert(len(body) - 1, element)


def clone_paragraph(document: Document, paragraph) -> None:
    insert_body_element(document, deepcopy(paragraph._p))


def replace_paragraph_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "필드 업데이트 필요"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_page_number_start(section_property, fmt: str | None = None, start: int | None = None) -> None:
    pg_num = section_property.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section_property.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def add_centered_title(document: Document, text: str, size: float = 16, before: float = 0, after: float = 12):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    set_east_asia_font(run, "HY신명조")
    return p


def add_body_paragraph(document: Document, text: str):
    p = document.add_paragraph(style="본문 내용")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.widow_control = True
    p.paragraph_format.keep_together = False
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_east_asia_font(run, "HY신명조")
    run.font.size = Pt(11)
    return p


def image_display_size(path: Path, max_width: float = 5.15, max_height: float = 6.55):
    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    width = max_width
    height = width / aspect
    if height > max_height:
        height = max_height
        width = height * aspect
    return Inches(width), Inches(height)


def add_caption(document: Document, caption: str, kind: str):
    label = "표" if kind == "table" else "그림"
    p = document.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = True
    p.add_run(f"[{label} ")
    add_field(p, f"SEQ {label} \\* ARABIC")
    p.add_run(f"] {caption}")
    for run in p.runs:
        set_east_asia_font(run, "HY신명조")
        run.font.size = Pt(9)
    return p


def add_figure(document: Document, filename: str, caption: str, kind: str = "figure"):
    path = FIGURE_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = document.add_paragraph(style="그림")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    width, height = image_display_size(path)
    p.add_run().add_picture(str(path), width=width, height=height)
    add_caption(document, caption, kind)


def add_equation(document: Document, equation: str):
    p = document.add_paragraph(style="수식")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run(equation)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)


def add_toc_page(document: Document, title: str, instruction: str):
    p = add_centered_title(document, title, size=16)
    p.paragraph_format.page_break_before = True
    field_p = document.add_paragraph()
    add_field(field_p, instruction)


def add_front_matter(document: Document, section_two_break) -> None:
    abstract_heading = document.add_paragraph("초    록", style="Heading 1")
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.paragraph_format.page_break_before = False
    abstract_heading.paragraph_format.space_before = Pt(6)
    abstract_heading.paragraph_format.space_after = Pt(18)
    abstract_heading.paragraph_format.keep_with_next = True
    for run in abstract_heading.runs:
        set_east_asia_font(run, "HY신명조")
        run.font.size = Pt(16)
    for text in KOREAN_ABSTRACT[:-1]:
        add_body_paragraph(document, text)
    keyword = document.add_paragraph()
    keyword.paragraph_format.space_before = Pt(12)
    keyword.add_run(KOREAN_ABSTRACT[-1])
    for run in keyword.runs:
        set_east_asia_font(run, "HY신명조")
        run.font.size = Pt(10.5)
    student = document.add_paragraph()
    student.add_run("학   번 : (추후 기입)")
    for run in student.runs:
        set_east_asia_font(run, "HY신명조")
        run.font.size = Pt(10.5)

    add_toc_page(document, "목   차", 'TOC \\o "1-2" \\h \\z \\u')
    add_toc_page(document, "표 목차", 'TOC \\h \\z \\c "표"')
    add_toc_page(document, "그림 목차", 'TOC \\h \\z \\c "그림"')
    insert_body_element(document, deepcopy(section_two_break))


def add_chapters(document: Document) -> None:
    for chapter_index, chapter in enumerate(CHAPTERS):
        heading = document.add_paragraph(chapter["title"], style="Heading 1")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.page_break_before = True
        heading.paragraph_format.keep_with_next = True
        for run in heading.runs:
            set_east_asia_font(run, "HY신명조")

        for section_index, section in enumerate(chapter["sections"]):
            section_heading = document.add_paragraph(section["title"], style="Heading 2")
            section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_heading.paragraph_format.keep_with_next = True
            if section_index > 0:
                section_heading.paragraph_format.page_break_before = True
            for run in section_heading.runs:
                set_east_asia_font(run, "HY신명조")

            for text in section["paragraphs"]:
                add_body_paragraph(document, text)
            for equation in section.get("equations", []):
                add_equation(document, equation)
            for filename, caption, kind in section.get("figures", []):
                add_figure(document, filename, caption, kind)


def add_references(document: Document) -> None:
    heading = document.add_paragraph("참고문헌", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True
    for reference in REFERENCES:
        reference = re.sub(r"^\[\d+\]\s*", "", reference)
        p = document.add_paragraph(style="References")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(reference)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)


def add_english_abstract(document: Document) -> None:
    heading = document.add_paragraph("Abstract", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True

    for line in ENGLISH_TITLE_LINES:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(15)

    byline = document.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.space_before = Pt(10)
    byline.paragraph_format.space_after = Pt(14)
    run = byline.add_run("Myung Jin\nDepartment of Electrical and Computer Engineering\nSeoul National University")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)

    for text in ENGLISH_ABSTRACT:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)


def normalize_sections(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(7.48125)
        section.page_height = Inches(10.2375)
    if len(document.sections) != 3:
        raise RuntimeError(f"Expected 3 sections, found {len(document.sections)}")

    s1, s2, s3 = document.sections
    s1.top_margin = Inches(0.98472)
    s1.bottom_margin = Inches(0.39375)
    s1.left_margin = Inches(1.18125)
    s1.right_margin = Inches(1.18125)
    s1.header_distance = Inches(0)
    s1.footer_distance = Inches(0)

    for section in (s2, s3):
        section.top_margin = Inches(0.7875)
        section.bottom_margin = Inches(0.59097)
        section.left_margin = Inches(1.18125)
        section.right_margin = Inches(1.18125)
        section.header_distance = Inches(0.59097)
        section.footer_distance = Inches(0.59097)

    # Preserve the template's distinct Roman and Arabic numbering systems.
    sect_prs = [section._sectPr for section in document.sections]
    set_page_number_start(sect_prs[1], fmt="lowerRoman", start=1)
    set_page_number_start(sect_prs[2], fmt="decimal", start=1)


def build() -> None:
    reference = Document(REFERENCE)
    cover_source = Document(TARGET_COVER)
    ref_paragraphs = list(reference.paragraphs)
    cover_paragraphs = list(cover_source.paragraphs)
    first_section_break = deepcopy(ref_paragraphs[40]._p)
    second_section_break = deepcopy(ref_paragraphs[117]._p)

    clear_body_keep_final_sectpr(reference)
    for paragraph in cover_paragraphs[:40]:
        clone_paragraph(reference, paragraph)
    insert_body_element(reference, first_section_break)

    cover_map = {
        0: "공학석사 학위논문",
        2: KOREAN_TITLE_LINES[0],
        3: KOREAN_TITLE_LINES[1],
        4: KOREAN_TITLE_LINES[2],
        6: ENGLISH_TITLE_LINES[0],
        7: ENGLISH_TITLE_LINES[1],
        8: ENGLISH_TITLE_LINES[2],
        11: " 2027 년 2 월",
        14: "서울대학교 대학원",
        15: "전기·정보공학부",
        16: "진 명",
        17: KOREAN_TITLE_LINES[0],
        18: KOREAN_TITLE_LINES[1],
        19: KOREAN_TITLE_LINES[2],
        21: ENGLISH_TITLE_LINES[0],
        22: ENGLISH_TITLE_LINES[1],
        23: ENGLISH_TITLE_LINES[2],
        25: "지도 교수  신 형 철",
        27: "이 논문을 공학석사 학위논문으로 제출함",
        28: "2027 년  2 월",
        30: "서울대학교 대학원",
        31: "전기·정보공학부",
        32: "진 명",
        34: "진명의 공학석사 학위논문을 인준함",
        35: "2027 년  2 월",
        37: "위 원 장         O O O         (인)",
        38: "부위원장         신 형 철         (인)",
        39: "위    원         O O O         (인)",
    }
    for index, text in cover_map.items():
        replace_paragraph_text(reference.paragraphs[index], text)
    # The template uses direct outline levels for cover-title typography.
    # Remove only the outline metadata so cover text does not pollute the TOC.
    for paragraph in reference.paragraphs[:40]:
        outline = paragraph._p.pPr.find(qn("w:outlineLvl"))
        if outline is not None:
            paragraph._p.pPr.remove(outline)

    add_front_matter(reference, second_section_break)
    add_chapters(reference)
    add_references(reference)
    add_english_abstract(reference)
    normalize_sections(reference)
    set_update_fields_on_open(reference)

    reference.core_properties.title = "극저온 3차원 낸드 플래시 메모리의 시간 단계 독립형 물리 기반 리텐션 손실 모델링"
    reference.core_properties.subject = "서울대학교 공학석사 학위논문"
    reference.core_properties.author = "진명"
    reference.core_properties.keywords = "3-D NAND Flash, cryogenic retention, TBT, attempt-to-escape frequency, time-step-free"
    reference.core_properties.comments = "양식 기준: 석사학위논문_이찬_초안.docx; 그림 출처: 로컬 논문 PDF 및 연구정리 PPTX"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    reference.save(OUTPUT)
    print(OUTPUT)
    print(f"paragraphs={len(reference.paragraphs)} sections={len(reference.sections)} inline_shapes={len(reference.inline_shapes)}")


if __name__ == "__main__":
    build()
