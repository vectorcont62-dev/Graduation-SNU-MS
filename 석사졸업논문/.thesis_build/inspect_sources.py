from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / ".thesis_build"
TEXT_OUT = OUT / "source_text"
PPT_MEDIA_OUT = OUT / "ppt_media"
TEXT_OUT.mkdir(parents=True, exist_ok=True)
PPT_MEDIA_OUT.mkdir(parents=True, exist_ok=True)

NS_P = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
NS_REL = {
    "r": "http://schemas.openxmlformats.org/package/2006/relationships"
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for idx, p in enumerate(doc.paragraphs):
        paragraphs.append(
            {
                "index": idx,
                "style": p.style.name if p.style else "",
                "text": p.text,
            }
        )
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append({"index": t_idx, "rows": rows})
    data = {
        "path": str(path),
        "sha256": sha256(path),
        "sections": len(doc.sections),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }
    (TEXT_OUT / f"{path.stem}.json").write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    with (TEXT_OUT / f"{path.stem}.txt").open("w", encoding="utf-8") as stream:
        for p in paragraphs:
            stream.write(
                f"[{p['index']:04d}] <{p['style']}> {p['text'].strip()}\n"
            )
        for table in tables:
            stream.write(f"\n[TABLE {table['index']}]\n")
            for row in table["rows"]:
                stream.write(" | ".join(cell.replace("\n", " / ") for cell in row) + "\n")
    return data


def natural_slide_key(name: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", name)
    return int(match.group(1)) if match else 0


def inspect_pptx(path: Path) -> dict:
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            (
                name
                for name in zf.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            ),
            key=natural_slide_key,
        )
        media_names = sorted(
            name for name in zf.namelist() if name.startswith("ppt/media/")
        )
        for name in media_names:
            target = PPT_MEDIA_OUT / Path(name).name
            with zf.open(name) as source, target.open("wb") as dest:
                shutil.copyfileobj(source, dest)

        slides = []
        for slide_name in slide_names:
            xml = etree.fromstring(zf.read(slide_name))
            texts = [
                node.text or ""
                for node in xml.xpath(".//a:t", namespaces=NS_P)
            ]
            rel_name = (
                "ppt/slides/_rels/"
                + Path(slide_name).name
                + ".rels"
            )
            rels = []
            if rel_name in zf.namelist():
                rel_xml = etree.fromstring(zf.read(rel_name))
                for rel in rel_xml.xpath(".//r:Relationship", namespaces=NS_REL):
                    rels.append(
                        {
                            "id": rel.get("Id"),
                            "type": rel.get("Type"),
                            "target": rel.get("Target"),
                        }
                    )
            slides.append(
                {
                    "slide": natural_slide_key(slide_name),
                    "text": "\n".join(t.strip() for t in texts if t.strip()),
                    "relationships": rels,
                }
            )
    data = {
        "path": str(path),
        "sha256": sha256(path),
        "slide_count": len(slides),
        "media_count": len(media_names),
        "slides": slides,
        "media": [Path(name).name for name in media_names],
    }
    (TEXT_OUT / f"{path.stem}.json").write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    with (TEXT_OUT / f"{path.stem}.txt").open("w", encoding="utf-8") as stream:
        for slide in slides:
            stream.write(f"\n===== SLIDE {slide['slide']} =====\n")
            stream.write(slide["text"] + "\n")
            media = [
                r["target"]
                for r in slide["relationships"]
                if r.get("target", "").startswith("../media/")
            ]
            if media:
                stream.write("[MEDIA] " + ", ".join(media) + "\n")
    return data


def inspect_pdf(path: Path) -> dict:
    pages = []
    captions = []
    with pdfplumber.open(path) as pdf:
        metadata = pdf.metadata or {}
        for number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
            pages.append({"page": number, "text": text})
            for line in text.splitlines():
                normalized = " ".join(line.split())
                if re.match(
                    r"^(Fig(?:ure)?\.?|Table)\s*\d+",
                    normalized,
                    flags=re.IGNORECASE,
                ):
                    captions.append({"page": number, "caption": normalized})
    data = {
        "path": str(path),
        "sha256": sha256(path),
        "page_count": len(pages),
        "metadata": metadata,
        "captions": captions,
    }
    (TEXT_OUT / f"{path.stem}.json").write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    with (TEXT_OUT / f"{path.stem}.txt").open("w", encoding="utf-8") as stream:
        for page in pages:
            stream.write(f"\n===== PAGE {page['page']} =====\n")
            stream.write(page["text"] + "\n")
    return data


def main() -> None:
    summary = {"docx": [], "pptx": [], "pdf": []}
    for path in sorted(ROOT.glob("*.docx")):
        summary["docx"].append(inspect_docx(path))
    for path in sorted(ROOT.glob("*.pptx")):
        summary["pptx"].append(inspect_pptx(path))
    for path in sorted(ROOT.glob("*.pdf")):
        summary["pdf"].append(inspect_pdf(path))

    compact = {
        kind: [
            {
                k: v
                for k, v in item.items()
                if k
                in {
                    "path",
                    "sha256",
                    "sections",
                    "paragraph_count",
                    "table_count",
                    "slide_count",
                    "media_count",
                    "page_count",
                    "captions",
                }
            }
            for item in items
        ]
        for kind, items in summary.items()
    }
    (OUT / "source_inventory.json").write_text(
        json.dumps(compact, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(compact, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
